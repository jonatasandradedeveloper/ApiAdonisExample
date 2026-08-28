#!/usr/bin/env python3
"""Gera o documento Word ÚNICO da Ana Karina (capa + sumário + todas as seções)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from generate_complete_docs import (
    ACCENT,
    GRAY,
    NAVY,
    add_footer,
    add_runs,
    first_h1,
    markdown_to_doc,
    set_cell_shading,
    set_run_font,
    style_doc,
)

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word"
ANA_DIR = ROOT / "ana"

OUTPUT = OUT / "Ana-Karina-Documento-Unico-People-Lead.docx"

ANA_SECTIONS = [
    ("00-ficha.md", "1. Ficha e contexto"),
    ("07-resumo-call-1.md", "2. Resumo da 1:1 #1 (28/07/2026)"),
    ("10-resumo-call-2.md", "3. Resumo da 1:1 #2 (28/08/2026)"),
    ("01-historico-1-1.md", "4. Histórico completo de 1:1s"),
    ("02-prioridades.md", "5. Prioridades e parecer do People Lead"),
    ("03-plano-carreira.md", "6. Plano de carreira"),
    ("06-planejamento-carreira-empresa.md", "7. Plano de carreira × o que a empresa precisa"),
    ("04-insumos-ciclo.md", "8. Insumos para o December Cycle"),
    ("05-acoes.md", "9. Ações e checklist"),
    ("11-dados-oficiais-rh-ciclo.md", "10. Dados oficiais RH e prazos"),
    ("08-workday-prioridades-abcd.md", "11. Workday — prioridades e ABCD"),
    ("09-autorreflexao-passo-a-passo.md", "12. Autorreflexão ABCD — passo a passo (Ana)"),
    ("12-abcd-form-people-lead-rascunho.md", "13. ABCD Form — rascunho do People Lead"),
    ("13-orientacao-ciclo-ana.md", "14. Orientação do ciclo (Ana Karina)"),
]

VEREDITO = (
    "Associate no BANCO BRADESCO (adm. 01/04/2026) com entrega em homologação, "
    "Copilot/agentes de IA como diferencial, inglês C1+ e AZ-900. Swift e AI-900 → FY27. "
    "Autorreflexão ABCD: ainda não submetida (28/08). Input PL até 11/09/2026."
)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("AVANADE · PEOPLE LEAD · DOCUMENTO ÚNICO")
    set_run_font(r, size=12, bold=True, color=NAVY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Ana Karina Caetano Dos Santos Marques")
    set_run_font(r, size=26, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Acompanhamento completo · carreira · prioridades · Workday · December Cycle FY26"
    )
    set_run_font(r, size=12, italic=True, color=GRAY)

    doc.add_paragraph()

    info = doc.add_table(rows=7, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Cargo", "Associate, Mobile & Device Dev · 12-Associate"),
        ("Cliente / projeto", "BANCO BRADESCO — app nativo (Kotlin / Android)"),
        ("Gestor projeto", "Rafael Coloda — rafael.coloda@avanade.com"),
        ("People Lead", "Jônatas Andrade Da Silva"),
        ("Admissão", "01/04/2026 · Alocação até 01/04/2027"),
        ("Prazos", "Ana: autorreflexão ABCD 28/08 (não submetida) · PL: 11/09"),
        ("Última atualização", "28 de agosto de 2026"),
    ]
    for i, (k, v) in enumerate(rows):
        pk = info.rows[i].cells[0].paragraphs[0]
        pv = info.rows[i].cells[1].paragraphs[0]
        add_runs(pk, k, bold=True)
        add_runs(pv, v)
        set_cell_shading(info.rows[i].cells[0], "F3F6FA")

    doc.add_paragraph()
    box = doc.add_paragraph()
    r = box.add_run("Síntese do People Lead")
    set_run_font(r, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    add_runs(p, VEREDITO)

    note = doc.add_paragraph()
    r = note.add_run(
        "Este arquivo reúne TUDO sobre a Ana Karina: fichas, calls, histórico, prioridades, "
        "plano de carreira, Workday, autorreflexão dela e rascunho da sua avaliação (PL). "
        "Documento vivo — atualizar após cada 1:1."
    )
    set_run_font(r, size=10, italic=True, color=GRAY)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Sumário", level=1)
    for _, title in ANA_SECTIONS:
        p = doc.add_paragraph()
        add_runs(p, title)
        p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    add_runs(p, "14. Orientação do ciclo (Ana Karina)")
    p.paragraph_format.space_after = Pt(6)

    tip = doc.add_paragraph()
    r = tip.add_run(
        "Dica: no Word, use Referências → Sumário para índice com números de página."
    )
    set_run_font(r, size=9, italic=True, color=GRAY)
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    style_doc(doc)
    add_footer(doc, "Ana Karina")
    add_cover(doc)
    add_toc(doc)

    for idx, (filename, section_title) in enumerate(ANA_SECTIONS):
        path = ANA_DIR / filename
        if not path.exists():
            print(f"SKIP  {filename} (not found)")
            continue
        if idx > 0:
            doc.add_page_break()
        md = path.read_text(encoding="utf-8")
        doc.add_heading(section_title, level=1)
        markdown_to_doc(md, doc, skip_first_h1=True, demote=1)

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def main() -> None:
    path = build()
    print(f"OK  {path.relative_to(ROOT)}  ({path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
