#!/usr/bin/env python3
"""Gera o documento Word COMPLETO de cada mentorada (capa + sumário + dossiê)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word"

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
UL_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
OL_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
CHECK_RE = re.compile(r"^(\s*)[-*]\s+\[([ xX])\]\s+(.*)$")
HR_RE = re.compile(r"^-{3,}\s*$")
FENCE_RE = re.compile(r"^```")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
BOLD_ITALIC = re.compile(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = "1F3A5F"

PEOPLE = {
    "andressa": {
        "nome": "Andressa Silva",
        "cargo": "Analyst, Back-End Developer",
        "cliente": "MAPFRE (sustentação / stack legado)",
        "arquivo": "Andressa-Silva-Historico-Completo-People-Lead.docx",
        "veredito": (
            "Prioridades com sentido de carreira, incompletas para o ciclo. "
            "Falta performance na MAPFRE. AI-900 está bloqueada (inglês + tentativas). "
            "Women In Tech mantém. Fechar FY26 com 3 frentes: MAPFRE + trilha técnica + Women In Tech. "
            "Promoção não é a tese deste FY. Não confundir com Ana Karina (BANCO BRADESCO)."
        ),
    },
    "ana": {
        "nome": "Ana Karina",
        "cargo": "Associate, Mobile & Device Dev",
        "cliente": "BANCO BRADESCO — app nativo (Kotlin / Android)",
        "arquivo": "Ana-Karina-Historico-Completo-People-Lead.docx",
        "veredito": (
            "Direção boa, portfólio inchado (7 prioridades). "
            "Copilot aplicado no Bradesco é o ouro. Inglês B2→C1+ já entregue. Swift é FY27. "
            "Recortar para 3 frentes: app + Copilot/IA + desenvolvimento do Associate. "
            "Promoção cedo demais salvo evidência excepcional. Não confundir com Andressa Silva (MAPFRE)."
        ),
    },
}

SECTION_ORDER = [
    ("00-ficha.md", "1. Ficha e contexto"),
    ("01-historico-1-1.md", "2. Histórico de 1:1s"),
    ("02-prioridades.md", "3. Prioridades e parecer"),
    ("03-plano-carreira.md", "4. Plano de carreira"),
    ("04-insumos-ciclo.md", "5. Insumos do December Cycle"),
    ("05-acoes.md", "6. Ações e checklist"),
]


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, *, size=11, bold=False, italic=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_runs(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    if not text:
        return
    parts = BOLD_ITALIC.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("***") and part.endswith("***"):
            run.text = part[3:-3]
            set_run_font(run, bold=True, italic=True)
        elif part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            set_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            set_run_font(run, size=9, name="Consolas", color=RGBColor(0x33, 0x33, 0x33))
        else:
            run.text = part
            set_run_font(run, bold=bold, italic=italic)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.color.rgb = NAVY
        h.font.name = "Calibri"


def add_footer(doc: Document, person_name: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"People Lead · {person_name} · documento vivo · atualizar a cada 1:1")
    set_run_font(run, size=8, color=GRAY, italic=True)


def add_cover(doc: Document, meta: dict) -> None:
    for _ in range(3):
        doc.add_paragraph()

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("AVANADE · PEOPLE LEAD")
    set_run_font(r, size=12, bold=True, color=NAVY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"Histórico completo — {meta['nome']}")
    set_run_font(r, size=28, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Acompanhamento de carreira · prioridades · December Cycle FY26")
    set_run_font(r, size=12, italic=True, color=GRAY)

    doc.add_paragraph()

    info = doc.add_table(rows=5, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Cargo", meta["cargo"]),
        ("Cliente / projeto", meta["cliente"]),
        ("People Lead", "Jônatas"),
        ("1ª conversa", "28 de julho de 2026"),
        ("Última atualização do dossiê", "24 de agosto de 2026"),
    ]
    for i, (k, v) in enumerate(rows):
        info.rows[i].cells[0].text = ""
        info.rows[i].cells[1].text = ""
        pk = info.rows[i].cells[0].paragraphs[0]
        pv = info.rows[i].cells[1].paragraphs[0]
        add_runs(pk, k, bold=True)
        add_runs(pv, v)
        set_cell_shading(info.rows[i].cells[0], "F3F6FA")

    doc.add_paragraph()
    box = doc.add_paragraph()
    r = box.add_run("Parecer do People Lead (síntese)")
    set_run_font(r, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    add_runs(p, meta["veredito"])

    note = doc.add_paragraph()
    r = note.add_run(
        "Este é um documento vivo. Depois de cada 1:1, atualize o histórico, "
        "o status das prioridades e o checklist de ações. Não reescreva o passado: acrescente."
    )
    set_run_font(r, size=10, italic=True, color=GRAY)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Sumário", level=1)
    items = [
        "1. Ficha e contexto",
        "2. Histórico de 1:1s",
        "3. Prioridades e parecer",
        "4. Plano de carreira",
        "5. Insumos do December Cycle",
        "6. Ações e checklist",
        "Anexo A — Orientação do ciclo (CALL geral)",
        "Anexo B — Templates (1:1, feedback, fechamento)",
    ]
    for item in items:
        p = doc.add_paragraph()
        add_runs(p, item)
        p.paragraph_format.space_after = Pt(6)

    tip = doc.add_paragraph()
    r = tip.add_run(
        "Dica: no Word, use Referências → Sumário se quiser páginas clicáveis depois de salvar."
    )
    set_run_font(r, size=9, italic=True, color=GRAY)
    doc.add_page_break()


def first_h1(md_text: str) -> str:
    for line in md_text.splitlines():
        m = HEADING_RE.match(line)
        if m and len(m.group(1)) == 1:
            return m.group(2).strip()
    return "Seção"


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_block(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return "|" in lines[idx] and TABLE_SEP_RE.match(lines[idx + 1] or "") is not None


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, cell_text, bold=True)
        set_cell_shading(cell, ACCENT)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(header)):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F6FA")
    doc.add_paragraph()


def markdown_to_doc(md_text: str, doc: Document, *, skip_first_h1: bool = False, demote: int = 0) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    first_h1_seen = False

    while i < len(lines):
        line = lines[i]

        if FENCE_RE.match(line):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, size=9, name="Consolas")
                p.paragraph_format.left_indent = Cm(0.5)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if HR_RE.match(line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and skip_first_h1 and not first_h1_seen:
                first_h1_seen = True
                i += 1
                continue
            style_level = min(max(level + demote, 1), 3)
            doc.add_heading(text, level=style_level)
            i += 1
            continue

        if is_table_block(lines, i):
            header = parse_table_row(lines[i])
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip() and not HEADING_RE.match(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        qm = QUOTE_RE.match(line)
        if qm:
            quote_lines = [qm.group(1)]
            i += 1
            while i < len(lines):
                qm2 = QUOTE_RE.match(lines[i])
                if not qm2:
                    break
                quote_lines.append(qm2.group(1))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            add_runs(p, " ".join(quote_lines), italic=True)
            for run in p.runs:
                run.font.color.rgb = GRAY
            continue

        cm = CHECK_RE.match(line)
        if cm:
            checked = cm.group(2).lower() == "x"
            text = cm.group(3)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, ("☑ " if checked else "☐ ") + text)
            i += 1
            continue

        um = UL_RE.match(line)
        if um:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, um.group(2))
            i += 1
            continue

        om = OL_RE.match(line)
        if om:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, om.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, line.strip())
        i += 1


def build_complete(person_key: str) -> Path:
    meta = PEOPLE[person_key]
    person_dir = ROOT / person_key
    dest = OUT / meta["arquivo"]

    doc = Document()
    style_doc(doc)
    add_footer(doc, meta["nome"])
    add_cover(doc, meta)
    add_toc(doc)

    for idx, (filename, section_title) in enumerate(SECTION_ORDER):
        path = person_dir / filename
        if not path.exists():
            continue
        if idx > 0:
            doc.add_page_break()
        md = path.read_text(encoding="utf-8")
        doc.add_heading(section_title, level=1)
        # Keep original H2/H3 under the numbered section
        markdown_to_doc(md, doc, skip_first_h1=True, demote=1)

    # Anexo A — ciclo
    doc.add_page_break()
    doc.add_heading("Anexo A — Orientação do ciclo (CALL geral)", level=1)
    ciclo = (ROOT / "01-ciclo-dezembro-orientacao.md").read_text(encoding="utf-8")
    markdown_to_doc(ciclo, doc, skip_first_h1=True, demote=1)

    # Anexo B — templates
    doc.add_page_break()
    doc.add_heading("Anexo B — Templates", level=1)
    intro = doc.add_paragraph()
    add_runs(
        intro,
        "Copiar os blocos abaixo para o histórico depois de cada 1:1, "
        "para coletar feedback de revisores e para fechar cada prioridade no Workday.",
    )
    for tpl in ["1-1.md", "coleta-feedback.md", "fechamento-prioridade.md"]:
        path = ROOT / "_templates" / tpl
        if not path.exists():
            continue
        md = path.read_text(encoding="utf-8")
        doc.add_heading(first_h1(md), level=2)
        markdown_to_doc(md, doc, skip_first_h1=True, demote=1)

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    return dest


def main() -> None:
    for key in PEOPLE:
        path = build_complete(key)
        print(f"OK  {path.relative_to(ROOT)}  ({path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
