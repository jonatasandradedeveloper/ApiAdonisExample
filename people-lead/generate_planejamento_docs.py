#!/usr/bin/env python3
"""Gera Word separados de planejamento de carreira × empresa."""

from __future__ import annotations

import importlib.util
from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word"

# Reuse markdown renderer from generate_complete_docs
spec = importlib.util.spec_from_file_location("gcd", ROOT / "generate_complete_docs.py")
gcd = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(gcd)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


PLANS = [
    {
        "src": ROOT / "andressa" / "06-planejamento-carreira-empresa.md",
        "dest": OUT / "Andressa-Silva-Planejamento-Carreira-e-Empresa.docx",
        "nome": "Andressa Silva",
        "cliente": "MAPFRE",
        "cargo": "Analyst, Back-End Developer",
    },
    {
        "src": ROOT / "ana" / "06-planejamento-carreira-empresa.md",
        "dest": OUT / "Ana-Karina-Planejamento-Carreira-e-Empresa.docx",
        "nome": "Ana Karina",
        "cliente": "BANCO BRADESCO",
        "cargo": "Associate, Mobile & Device Dev",
    },
]


def build(plan: dict) -> Path:
    md = plan["src"].read_text(encoding="utf-8")
    doc = Document()
    gcd.style_doc(doc)
    gcd.add_footer(doc, plan["nome"])

    # cover
    for _ in range(2):
        doc.add_paragraph()
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("AVANADE · PEOPLE LEAD · PLANEJAMENTO")
    gcd.set_run_font(r, size=12, bold=True, color=gcd.NAVY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"Plano de carreira × o que a empresa precisa")
    gcd.set_run_font(r, size=22, bold=True, color=gcd.NAVY)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run(plan["nome"])
    gcd.set_run_font(r, size=18, bold=True, color=gcd.NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{plan['cargo']} · {plan['cliente']}")
    gcd.set_run_font(r, size=12, italic=True, color=gcd.GRAY)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Documento separado · editável · revisar a cada 1:1")
    gcd.set_run_font(r, size=10, italic=True, color=gcd.GRAY)

    doc.add_paragraph()
    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Pessoa", plan["nome"]),
        ("Cliente", plan["cliente"]),
        ("People Lead", "Jônatas"),
        ("Atualizado em", "24/08/2026"),
    ]
    for i, (k, v) in enumerate(rows):
        info.rows[i].cells[0].text = ""
        info.rows[i].cells[1].text = ""
        gcd.add_runs(info.rows[i].cells[0].paragraphs[0], k, bold=True)
        gcd.add_runs(info.rows[i].cells[1].paragraphs[0], v)
        gcd.set_cell_shading(info.rows[i].cells[0], "F3F6FA")

    doc.add_page_break()
    gcd.markdown_to_doc(md, doc, skip_first_h1=True)

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(plan["dest"])
    return plan["dest"]


def main() -> None:
    for plan in PLANS:
        path = build(plan)
        print(f"OK  {path.relative_to(ROOT)}  ({path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
