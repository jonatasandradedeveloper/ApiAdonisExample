#!/usr/bin/env python3
"""Convert people-lead markdown dossiers into editable Word (.docx) files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word"

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
UL_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
OL_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
CHECK_RE = re.compile(r"^(\s*)[-*]\s+\[([ xX])\]\s+(.*)$")
HR_RE = re.compile(r"^-{3,}\s*$")
FENCE_RE = re.compile(r"^```")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
BOLD_ITALIC = re.compile(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_runs(paragraph, text: str, *, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    if not text:
        return
    parts = BOLD_ITALIC.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("***") and part.endswith("***"):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run.text = part
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for i in range(1, 4):
        h = styles[f"Heading {i}"]
        h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        h.font.name = "Calibri"


def add_title_block(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    if subtitle:
        s = doc.add_paragraph()
        r = s.add_run(subtitle)
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def is_table_block(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return "|" in lines[idx] and TABLE_SEP_RE.match(lines[idx + 1] or "") is not None


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, cell_text, bold=True)
        set_cell_shading(cell, "1F3A5F")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(header)):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F6FA")
    doc.add_paragraph()


def markdown_to_doc(md_text: str, doc: Document, *, skip_first_h1: bool = False) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    first_h1_seen = False

    while i < len(lines):
        line = lines[i]

        if FENCE_RE.match(line):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(0.5)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if HR_RE.match(line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and skip_first_h1 and not first_h1_seen:
                first_h1_seen = True
                i += 1
                continue
            # Word heading styles only go reliably to 9; we use 1-3 mapped
            style_level = min(level, 3)
            doc.add_heading(text, level=style_level)
            i += 1
            continue

        if is_table_block(lines, i):
            header = parse_table_row(lines[i])
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip() and not HEADING_RE.match(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        qm = QUOTE_RE.match(line)
        if qm:
            quote_lines = [qm.group(1)]
            i += 1
            while i < len(lines):
                qm2 = QUOTE_RE.match(lines[i])
                if not qm2:
                    break
                quote_lines.append(qm2.group(1))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            add_runs(p, " ".join(quote_lines), italic=True)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        cm = CHECK_RE.match(line)
        if cm:
            checked = cm.group(2).lower() == "x"
            text = cm.group(3)
            p = doc.add_paragraph(style="List Bullet")
            mark = "☑ " if checked else "☐ "
            add_runs(p, mark + text)
            i += 1
            continue

        um = UL_RE.match(line)
        if um:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, um.group(2))
            i += 1
            continue

        om = OL_RE.match(line)
        if om:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, om.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, line.strip())
        i += 1


def first_h1(md_text: str) -> str:
    for line in md_text.splitlines():
        m = HEADING_RE.match(line)
        if m and len(m.group(1)) == 1:
            return m.group(2).strip()
    return "Documento"


def convert_file(src: Path, dest: Path, *, subtitle: str | None = None) -> None:
    md = src.read_text(encoding="utf-8")
    doc = Document()
    style_doc(doc)
    title = first_h1(md)
    add_title_block(doc, title, subtitle or "People Lead · documento editável · atualizar a cada 1:1")
    markdown_to_doc(md, doc, skip_first_h1=True)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    print(f"OK  {dest.relative_to(ROOT)}")


def build_dossier(person_dir: Path, dest: Path, person_name: str) -> None:
    order = [
        "00-ficha.md",
        "01-historico-1-1.md",
        "02-prioridades.md",
        "03-plano-carreira.md",
        "04-insumos-ciclo.md",
        "05-acoes.md",
    ]
    doc = Document()
    style_doc(doc)
    add_title_block(
        doc,
        f"Dossiê People Lead — {person_name}",
        "Histórico editável · prioridades · plano de carreira · December Cycle FY26",
    )
    for name in order:
        path = person_dir / name
        if not path.exists():
            continue
        md = path.read_text(encoding="utf-8")
        # page break between sections except first
        if name != order[0]:
            doc.add_page_break()
        section_title = first_h1(md)
        doc.add_heading(section_title, level=1)
        markdown_to_doc(md, doc, skip_first_h1=True)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    print(f"OK  {dest.relative_to(ROOT)}")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)

    singles = [
        (ROOT / "README.md", OUT / "00-README-kit.docx", "Índice do kit People Lead"),
        (ROOT / "00-dashboard.md", OUT / "00-dashboard.docx", "Visão mensal das duas mentoradas"),
        (ROOT / "01-ciclo-dezembro-orientacao.md", OUT / "01-ciclo-dezembro-orientacao.docx", "CALL geral → orientação do ciclo"),
        (ROOT / "02-parecer-prioridades.md", OUT / "02-parecer-prioridades.docx", "As prioridades fazem sentido?"),
    ]
    for src, dest, sub in singles:
        convert_file(src, dest, subtitle=sub)

    for person, label in [("andressa", "Andressa"), ("ana", "Ana")]:
        person_dir = ROOT / person
        person_out = OUT / person
        for md in sorted(person_dir.glob("*.md")):
            convert_file(
                md,
                person_out / (md.stem + ".docx"),
                subtitle=f"{label} · People Lead · documento editável",
            )
        build_dossier(
            person_dir,
            OUT / f"DOSSIE-{label.upper()}.docx",
            label,
        )

    # templates
    tpl_out = OUT / "_templates"
    for md in sorted((ROOT / "_templates").glob("*.md")):
        convert_file(md, tpl_out / (md.stem + ".docx"), subtitle="Template editável")

    print("\nDone. Word files in people-lead/word/")


if __name__ == "__main__":
    main()
